"""
Reusable question-bank ingestion core.

Extracts questions from a single PDF (Groq Vision, per page — works on scanned
and digital papers, crops diagrams) or a single .docx (python-docx text → Groq
text). Persists to QuestionBank deduped by question_hash, unverified
(is_verified=False) so everything lands in the admin review queue. A
SourceDocument (deduped by file hash) ties questions back to their file.

Both the single-file `ingest_question_bank` command and the bulk `ingest_folder`
command call into here, so the extraction logic lives in one place.
"""
import base64
import hashlib
import io
import json
import re
import time

import fitz  # PyMuPDF
from django.core.files.base import ContentFile

from ai_engine.models import (
    CaseStudyPart, Difficulty, IngestionStatus, MCQOption,
    QuestionBank, QuestionType, SourceDocument,
)

VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
TEXT_MODEL = "llama-3.3-70b-versatile"

# Local Ollama (text-only) fallback provider. Qwen2.5-Coder has no vision, so the
# Ollama path uses each file's TEXT (PDF text layer / docx) — scanned PDFs and
# embedded diagrams are not handled by this provider.
OLLAMA_BASE_URL = "http://host.docker.internal:11434"
OLLAMA_MODEL = "qwen2.5-coder:7b"


# ── LaTeX → plain-text cleaner ────────────────────────────────────────────────
_SUB_DIGITS = str.maketrans('0123456789', '₀₁₂₃₄₅₆₇₈₉')
_SUP_MAP = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'}
_SUB_LTRS = {'n': 'ₙ', 'm': 'ₘ', 'p': 'ₚ', 'k': 'ₖ', 'q': 'q', 'i': 'ᵢ', 'r': 'ᵣ'}
_FRAC_UNI = {
    ('1', '2'): '½', ('1', '3'): '⅓', ('2', '3'): '⅔',
    ('1', '4'): '¼', ('3', '4'): '¾',
    ('1', '8'): '⅛', ('3', '8'): '⅜', ('5', '8'): '⅝', ('7', '8'): '⅞',
}


def _sup(s):
    return ''.join(_SUP_MAP.get(c, c) for c in s)


def _sub(s):
    out = ''
    for c in s:
        if c.isdigit():
            out += c.translate(_SUB_DIGITS)
        elif c in _SUB_LTRS:
            out += _SUB_LTRS[c]
        elif c == '+':
            out += '₊'
        elif c == '-':
            out += '₋'
        else:
            out += c
    return out


def clean_latex_text(text: str) -> str:
    def _mixed(m):
        w, n, d = m.group(1), m.group(2), m.group(3)
        return w + _FRAC_UNI.get((n, d), f'{n}/{d}')
    text = re.sub(r'(-?\d+)\\frac\{(\d+)\}\{(\d+)\}', _mixed, text)

    def _frac(m):
        n, d = m.group(1).strip(), m.group(2).strip()
        return _FRAC_UNI.get((n, d), f'{n}/{d}') if (n.isdigit() and d.isdigit()) else f'{n}/{d}'
    text = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', _frac, text)

    def _sqrt(m):
        inner = m.group(1).strip()
        return ('√' + inner) if re.match(r'^[\w.]+$', inner) else ('√(' + inner + ')')
    text = re.sub(r'\\sqrt\{([^{}]+)\}', _sqrt, text)
    text = re.sub(r'\\sqrt\s*(\d+)', lambda m: '√' + m.group(1), text)
    text = re.sub(r'\^\{([^{}]+)\}', lambda m: _sup(m.group(1)), text)
    text = re.sub(r'\^(\d)', lambda m: _SUP_MAP.get(m.group(1), m.group(1)), text)
    text = re.sub(r'_\{([^{}]+)\}', lambda m: _sub(m.group(1)), text)
    text = re.sub(r'_(\d)', lambda m: m.group(1).translate(_SUB_DIGITS), text)
    text = re.sub(r'_([nmkpq])', lambda m: _SUB_LTRS.get(m.group(1), m.group(1)), text)
    for pat in [r'\\cdot\\cdot\\cdot', r'\\cdots', r'\\ldots', r'\\dots']:
        text = re.sub(pat, '...', text)
    for pat, rep in [
        (r'\\times', '×'), (r'\\div', '÷'), (r'\\pm', '±'), (r'\\mp', '∓'),
        (r'\\neq', '≠'), (r'\\ne\b', '≠'), (r'\\leq', '≤'), (r'\\geq', '≥'),
        (r'\\le\b', '≤'), (r'\\ge\b', '≥'), (r'\\infty', '∞'),
        (r'\\alpha', 'α'), (r'\\beta', 'β'), (r'\\theta', 'θ'), (r'\\pi', 'π'),
        (r'\\left', ''), (r'\\right', ''),
    ]:
        text = re.sub(pat, rep, text)
    text = text.replace('$', '')
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = text.replace('{', '').replace('}', '')
    text = re.sub(r' {2,}', ' ', text).strip()
    return text


# ── Prompts ───────────────────────────────────────────────────────────────────
VISION_PROMPT = """
You are an expert extracting and SOLVING questions from a question-paper page.

IMPORTANT: The page may NOT have answers marked. SOLVE each question yourself
to determine the correct option.

Return ONLY a valid JSON object — no markdown fences, compact single-line.

Schema:
{"questions":[{"number":1,"type":"MCQ","marks":1,"text":"full question text","has_diagram":true,"diagram_bbox":[0.25,0.10,0.75,0.38],"options":[{"label":"a","text":"opt","is_correct":false}],"answer_text":"brief working","difficulty":"medium","board_tag":"BOARD 2023","case_parts":[]}]}

Rules:
1. type ∈ MCQ, ASSERTION_REASON, VERY_SHORT, SHORT, LONG, CASE
2. diagram_bbox: if has_diagram, [x1,y1,x2,y2] as 0.0-1.0 fractions of the page (origin top-left); omit if none.
3. difficulty: easy/medium/hard.
4. board_tag: e.g. "BOARD 2023"; null if absent.
5. SOLVE each MCQ and mark exactly one option is_correct:true; put working in answer_text.
6. For CASE: fill case_parts and solve each sub-question.
7. Include ALL questions on the page; ignore headers/footers/page numbers.
8. NEVER use LaTeX — write plain Unicode (a₁, x², 1/2, √2, ×, ÷, ≤, ≥, π …). No $, \\frac, \\sqrt, backslashes, _{}, ^{}.
9. Each "text" must be a COMPLETE standalone question — never output an answer
   choice or single label as its own question (put those in "options").
"""


def build_text_prompt(subject: str, chapter: str, grade: str, board: str, content: str) -> str:
    return f"""Extract and SOLVE every question from this {board} {grade} {subject} document (chapter/topic: {chapter}).

Return ONLY a valid JSON object — no markdown, compact:
{{"questions":[{{"type":"MCQ","marks":1,"text":"full question text","options":[{{"label":"a","text":"opt","is_correct":false}}],"answer_text":"brief working/answer","difficulty":"medium"}}]}}

Rules:
1. type ∈ MCQ, ASSERTION_REASON, VERY_SHORT, SHORT, LONG, CASE.
2. SOLVE each MCQ/AR; mark exactly one option is_correct:true. Non-option types: options=[].
3. difficulty: easy/medium/hard.
4. Include every question; ignore headers/footers/instructions.
5. NEVER use LaTeX — plain Unicode only (a₁, x², 1/2, √2, ×, ÷, ≤, ≥, π …). No $, backslashes, _{{}}, ^{{}}.
6. Each "text" MUST be a COMPLETE, standalone question. NEVER output an answer
   choice, a single word, or a classification label (e.g. "bijective function")
   as its own question — those belong in "options" of their parent question.
7. If a question's math cannot be written cleanly in plain Unicode, OMIT that
   question entirely rather than emitting garbled/repeated characters.

===DOCUMENT===
{content}
===END==="""


# ── JSON recovery (handles truncated model output) ────────────────────────────
def _extract_partial_questions(text: str) -> list:
    arr_start = text.find('[', text.find('"questions"') if '"questions"' in text else 0)
    if arr_start == -1:
        return []
    questions = []
    i = arr_start + 1
    while i < len(text):
        while i < len(text) and text[i] in ' \n\r\t':
            i += 1
        if i >= len(text) or text[i] != '{':
            break
        depth, j, in_str, esc = 0, i, False, False
        while j < len(text):
            c = text[j]
            if esc:
                esc = False
            elif c == '\\' and in_str:
                esc = True
            elif c == '"':
                in_str = not in_str
            elif not in_str:
                if c == '{':
                    depth += 1
                elif c == '}':
                    depth -= 1
                    if depth == 0:
                        try:
                            q = json.loads(text[i:j + 1])
                            if "text" in q:
                                questions.append(q)
                        except json.JSONDecodeError:
                            pass
                        i = j + 1
                        break
            j += 1
        else:
            break
        while i < len(text) and text[i] in ' \n\r\t,':
            i += 1
    return questions


def safe_json(text: str):
    if not text:
        return None
    text = re.sub(r"^```[a-z]*\n?", "", text.strip())
    text = re.sub(r"\n?```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    questions = _extract_partial_questions(text)
    return {"questions": questions} if questions else None


def _qtype(raw: str) -> str:
    return {
        "MCQ": QuestionType.MCQ, "ASSERTION_REASON": QuestionType.ASSERTION_REASON,
        "VERY_SHORT": QuestionType.VERY_SHORT, "SHORT": QuestionType.SHORT,
        "LONG": QuestionType.LONG, "CASE": QuestionType.CASE, "REARRANGE": QuestionType.REARRANGE,
    }.get((raw or "").upper(), QuestionType.SHORT)


def _difficulty(marks: int, explicit: str) -> str:
    mapping = {"easy": Difficulty.EASY, "medium": Difficulty.MEDIUM, "hard": Difficulty.HARD}
    if explicit in mapping:
        return mapping[explicit]
    return {1: Difficulty.EASY, 2: Difficulty.MEDIUM}.get(marks, Difficulty.HARD)


def compute_hash(subject: str, chapter: str, text: str) -> str:
    return hashlib.sha256(f"{subject}{chapter}{text.strip().lower()}".encode()).hexdigest()


def _crop_image(pix, bbox, page_w, page_h) -> bytes:
    from PIL import Image
    x1 = max(0, int(bbox[0] * page_w)); y1 = max(0, int(bbox[1] * page_h))
    x2 = min(page_w, int(bbox[2] * page_w)); y2 = min(page_h, int(bbox[3] * page_h))
    if x2 <= x1 or y2 <= y1:
        return b""
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    buf = io.BytesIO()
    img.crop((x1, y1, x2, y2)).save(buf, format="PNG")
    return buf.getvalue()


def _file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# ── Quality guards (small local models emit garbled formulas + fragments) ─────
# Exotic unicode ranges: subscripts, superscripts, currency, combining marks.
def _is_exotic(ch):
    o = ord(ch)
    return (0x2080 <= o <= 0x20CF) or (0x2070 <= o <= 0x207F) or (0x0300 <= o <= 0x036F)


def _looks_garbled(text):
    """A long run of subscript/super/combining chars, or a high overall ratio,
    means a mangled formula (e.g. f(x)=… exploded into ₁₂₃...)."""
    run = 0
    for ch in text:
        if _is_exotic(ch):
            run += 1
            if run >= 6:
                return True
        else:
            run = 0
    nonspace = [c for c in text if not c.isspace()]
    return bool(nonspace) and sum(_is_exotic(c) for c in nonspace) / len(nonspace) > 0.25


_Q_CUES = ('what', 'which', 'find', 'prove', 'define', 'state', 'show', 'evaluate',
           'solve', 'calculate', 'how', 'why', 'is ', 'are ', 'let ', 'given',
           'if ', 'draw', 'write', 'explain', 'determine', 'express', 'verify')


def _is_fragment(text, qtype, options):
    """Stray answer-choice / label saved as its own 'question'. Only applies to
    option-less free-text types (a real MCQ stem keeps its options)."""
    if options or qtype.upper() not in ('SHORT', 'VERY_SHORT', ''):
        return False
    t = text.strip().lower()
    if len(t) < 20:
        return True
    # Short, with no question mark and none of the usual question/command cues →
    # almost certainly a stray label/phrase (e.g. "many one, into function.").
    return len(t) < 45 and '?' not in t and not any(cue in t for cue in _Q_CUES)


# ── Persist one extracted question ─────────────────────────────────────────────
def _save_question(q, *, subject, chapter, src_doc, source_page, pix, page_w, page_h, dry_run, log):
    """Returns 'created' | 'skipped' | 'error'."""
    text = clean_latex_text((q.get("text") or "").strip())
    if not text:
        return 'skipped'

    # Drop low-quality extractions (common with small local models) before they
    # pollute the review queue.
    qtype = (q.get("type") or "").upper()
    options = q.get("options") or []
    if _looks_garbled(text):
        log(f"      skipped (garbled formula): {text[:40]}…")
        return 'skipped'
    if _is_fragment(text, qtype, options):
        log(f"      skipped (fragment, not a question): {text[:40]}")
        return 'skipped'
    if qtype in ('MCQ', 'ASSERTION_REASON') and not options:
        log(f"      skipped (MCQ with no options): {text[:40]}")
        return 'skipped'

    q_hash = compute_hash(subject, chapter, text)

    if dry_run:
        correct = next((o.get("label") for o in q.get("options", []) if o.get("is_correct")), "?")
        log(f"      [{q.get('type', '?')} {q.get('marks', '?')}m] ans={correct}  {text[:70]}")
        return 'created'

    if QuestionBank.objects.filter(question_hash=q_hash).exists():
        return 'skipped'

    has_image = bool(q.get("has_diagram")) and pix is not None
    diagram_file = None
    image_description = ""
    if has_image and q.get("diagram_bbox"):
        crop = _crop_image(pix, q["diagram_bbox"], page_w, page_h)
        if crop:
            safe = re.sub(r"[^a-z0-9]", "_", text[:30].lower())
            diagram_file = ContentFile(crop, name=f"p{source_page}_{safe}.png")
            image_description = f"Diagram for: {text[:80]}"
        else:
            has_image = False
    else:
        has_image = False

    try:
        marks_raw = int(q.get("marks") or 1)
        qb = QuestionBank(
            subject=subject, chapter=chapter,
            question_type=_qtype(q.get("type", "SHORT")),
            marks=max(1, min(20, marks_raw)),
            difficulty=_difficulty(max(1, marks_raw), q.get("difficulty") or ""),
            question_text=text,
            answer_text=q.get("answer_text") or "",
            has_image=has_image and diagram_file is not None,
            image_description=image_description,
            tags=[q["board_tag"]] if q.get("board_tag") else [],
            question_hash=q_hash,
            is_ai_generated=False, is_verified=False,
            source_document=src_doc, source_page=source_page,
        )
        if diagram_file:
            qb.image.save(diagram_file.name, diagram_file, save=False)
        qb.save()

        for i, opt in enumerate((q.get("options") or [])[:5]):
            if not isinstance(opt, dict):
                continue
            MCQOption.objects.create(
                question=qb,
                option_label=str(opt.get("label", "")).strip()[:1].lower() or chr(97 + i),
                option_text=str(opt.get("text", "")).strip(),
                is_correct=bool(opt.get("is_correct")),
                order=i,
            )
        for part in q.get("case_parts") or []:
            if not isinstance(part, dict):
                continue
            CaseStudyPart.objects.get_or_create(
                parent_question=qb, part_number=int(part.get("part_number", 1)),
                defaults=dict(
                    part_text=str(part.get("text", "")).strip(),
                    part_answer=str(part.get("answer", "")).strip(),
                    question_type=_qtype(part.get("type", "SHORT")),
                    marks=max(1, int(part.get("marks") or 1)),
                    has_image=bool(part.get("has_diagram")),
                ),
            )
        return 'created'
    except Exception as exc:
        log(f"      DB error: {exc}")
        return 'error'


def _open_source_doc(path, meta, total_pages, dry_run):
    if dry_run:
        return None
    src_doc, created = SourceDocument.objects.get_or_create(
        file_hash=_file_hash(path),
        defaults={
            "title": f"{meta['subject']} — {meta['chapter']}", "subject": meta['subject'],
            "chapter": meta['chapter'], "board": meta['board'], "class_grade": meta['grade'],
            "page_count": total_pages, "status": IngestionStatus.PROCESSING,
        },
    )
    if not created:
        src_doc.status = IngestionStatus.PROCESSING
        src_doc.save(update_fields=["status"])
    return src_doc


def _call_vision(client, model_id, b64, log):
    for attempt in range(4):
        try:
            resp = client.chat.completions.create(
                model=model_id,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    {"type": "text", "text": VISION_PROMPT},
                ]}],
                max_tokens=8192, temperature=0.1,
            )
            return resp.choices[0].message.content
        except Exception as exc:
            err = str(exc)
            if "429" in err or "rate_limit" in err.lower():
                wait = 60 * (attempt + 1)
                log(f"      rate limited — waiting {wait}s (attempt {attempt + 1}/4)…")
                time.sleep(wait)
            else:
                log(f"      GROQ ERROR: {exc}")
                return None
    return None


def _call_ollama(base_url, model, prompt, log, num_ctx=8192, timeout=600):
    """Call a local Ollama text model via its native /api/chat (JSON-forced)."""
    import requests as rq
    try:
        r = rq.post(base_url.rstrip('/') + '/api/chat', json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False, "format": "json",
            "options": {"temperature": 0.1, "num_ctx": num_ctx},
        }, timeout=timeout)
        r.raise_for_status()
        return r.json().get("message", {}).get("content", "")
    except Exception as exc:
        log(f"      OLLAMA ERROR: {exc}")
        return None


# ── Public: ingest one PDF ──────────────────────────────────────────────────
def ingest_pdf(pdf_path, *, subject, chapter, board, grade, provider='groq', client=None,
               model_id=None, ollama_base_url=OLLAMA_BASE_URL, ollama_model=OLLAMA_MODEL,
               dpi=150, delay=5, dry_run=False, max_pages=0, log=print):
    from PIL import Image
    meta = {"subject": subject, "chapter": chapter, "board": board, "grade": grade}
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    src_doc = _open_source_doc(pdf_path, meta, total_pages, dry_run)
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    vision_model = model_id or VISION_MODEL
    created = skipped = errors = 0
    pages = list(range(total_pages))
    if max_pages:
        pages = pages[:max_pages]

    def _encode_crop(p, y0, y1):
        img = Image.frombytes("RGB", [p.width, p.height], p.samples).crop((0, y0, p.width, y1))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.standard_b64encode(buf.getvalue()).decode()

    for page_idx in pages:
        page = doc[page_idx]
        page_num = page_idx + 1
        if not dry_run and src_doc:
            already = QuestionBank.objects.filter(source_document=src_doc, source_page=page_num).count()
            if already > 0:
                log(f"    page {page_num}/{total_pages} — already saved ({already}) — skipped")
                skipped += already
                continue

        pix = None
        page_w = page_h = 0
        if provider == 'ollama':
            # Text-only: use the page's text layer (no vision, no diagram crop).
            ptext = page.get_text()
            if not ptext.strip():
                log(f"    page {page_num}/{total_pages} — no text layer (scanned) — needs a vision provider, skipped")
                errors += 1
                continue
            raw = _call_ollama(ollama_base_url, ollama_model,
                               build_text_prompt(subject, chapter, grade, board, ptext[:12000]), log)
            parsed = safe_json(raw)
        else:
            pix = page.get_pixmap(matrix=matrix)
            page_w, page_h = pix.width, pix.height
            parsed = safe_json(_call_vision(client, vision_model, _encode_crop(pix, 0, page_h), log))
            if not parsed or "questions" not in parsed:
                mid = page_h // 2
                collected = []
                for y0, y1 in [(0, mid), (mid, page_h)]:
                    time.sleep(delay)
                    half = safe_json(_call_vision(client, vision_model, _encode_crop(pix, y0, y1), log))
                    if half and "questions" in half:
                        collected.extend(half["questions"])
                parsed = {"questions": collected} if collected else None

        if not parsed or "questions" not in parsed:
            log(f"    page {page_num}/{total_pages} — extraction failed, skipped")
            errors += 1
            continue

        qs = parsed["questions"]
        log(f"    page {page_num}/{total_pages} — {len(qs)} question(s)")
        for q in qs:
            r = _save_question(q, subject=subject, chapter=chapter, src_doc=src_doc,
                               source_page=page_num, pix=pix, page_w=page_w, page_h=page_h,
                               dry_run=dry_run, log=log)
            created += (r == 'created'); skipped += (r == 'skipped'); errors += (r == 'error')
        if provider == 'groq' and page_idx < pages[-1]:
            time.sleep(delay)

    if src_doc and not dry_run:
        src_doc.status = IngestionStatus.DONE if errors == 0 else IngestionStatus.FAILED
        src_doc.save(update_fields=["status"])
    doc.close()
    return {"created": created, "skipped": skipped, "errors": errors, "pages": total_pages}


# ── Public: ingest one DOCX (text only — figures not captured) ────────────────
def ingest_docx(docx_path, *, subject, chapter, board, grade, provider='groq', client=None,
                model_id=None, ollama_base_url=OLLAMA_BASE_URL, ollama_model=OLLAMA_MODEL,
                dry_run=False, log=print):
    from docx import Document
    document = Document(docx_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n".join(parts)[:14000]
    if not content.strip():
        log("    (no extractable text) — skipped")
        return {"created": 0, "skipped": 0, "errors": 0, "pages": 0}

    meta = {"subject": subject, "chapter": chapter, "board": board, "grade": grade}
    src_doc = _open_source_doc(docx_path, meta, None, dry_run)
    prompt = build_text_prompt(subject, chapter, grade, board, content)

    created = skipped = errors = 0
    try:
        if provider == 'ollama':
            parsed = safe_json(_call_ollama(ollama_base_url, ollama_model, prompt, log))
        else:
            resp = client.chat.completions.create(
                model=model_id or TEXT_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1, max_tokens=8000,
            )
            parsed = safe_json(resp.choices[0].message.content)
    except Exception as exc:
        log(f"    LLM ERROR: {exc}")
        if src_doc:
            src_doc.status = IngestionStatus.FAILED
            src_doc.save(update_fields=["status"])
        return {"created": 0, "skipped": 0, "errors": 1, "pages": 0}

    qs = (parsed or {}).get("questions", [])
    log(f"    {len(qs)} question(s) extracted")
    for q in qs:
        r = _save_question(q, subject=subject, chapter=chapter, src_doc=src_doc,
                           source_page=None, pix=None, page_w=0, page_h=0,
                           dry_run=dry_run, log=log)
        created += (r == 'created'); skipped += (r == 'skipped'); errors += (r == 'error')

    if src_doc and not dry_run:
        src_doc.status = IngestionStatus.DONE if errors == 0 else IngestionStatus.FAILED
        src_doc.save(update_fields=["status"])
    return {"created": created, "skipped": skipped, "errors": errors, "pages": 0}
